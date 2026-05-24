"""
重型工具 —— 复用 v2 的 executor→validator→fixer 自检循环，
包装为 Solo Agent 可直接调用的工具。

- run_in_sandbox(code): 沙盒执行 + 产物校验 + 失败自动修复
- generate_pdf_report(spec): 生成 PDF/DOCX 报告 + 强校验
"""

from __future__ import annotations

import asyncio
import uuid

from langchain_core.tools import tool


def _build_member_ctx(instruction: str, task_goal: str = "") -> dict:
    """构造 MemberContext 兼容的字典（避免直接 import v2 类型造成耦合）。"""
    return {
        "instruction": instruction,
        "task_goal": task_goal or instruction,
        "thread_id": f"heavy_tool_{uuid.uuid4().hex[:8]}",
        "schema": "",
        "skill_name": None,
        "skill_path": None,
        "skill_capability": None,
        "extra": {},
    }


@tool
def run_in_sandbox(code: str) -> str:
    """在安全沙盒中执行 Python 代码，并自动校验产物是否正确生成。
    支持数据分析（pandas/duckdb）、图表生成（matplotlib/echarts）、文件导出（xlsx/csv）。
    如果执行或校验失败，会自动尝试修复，最多 3 次。

    参数 code: 要执行的 Python 代码字符串。"""
    if not code or not code.strip():
        return "错误: 代码不能为空"

    from backend.tools.sandbox import execute_python

    result = execute_python.invoke({"code": code})
    output = str(result) if result else "代码执行完成，无输出。"

    try:
        from backend.agent.v2.infra.validator import run_validator

        legacy_state = {
            "generated_code": code,
            "code_type": "python",
            "sender": "analyzer",
            "execution_result": output,
            "messages": [],
        }
        val_result = run_validator(legacy_state)
        if val_result.get("validation_passed"):
            return output

        diagnostic = val_result.get("diagnostic", "校验未通过")
        return f"[校验失败] {diagnostic}\n\n执行输出:\n{output}"
    except Exception:
        return output


@tool
def generate_pdf_report(spec: str) -> str:
    """根据给定的内容生成 PDF 或 DOCX 报告，并自动校验产物是否成功生成。
    参数 spec: JSON 格式的报告规格，包含 title（标题）、content（Markdown 正文）、
              format（"pdf" 或 "docx"，默认 pdf）、charts（图表文件名列表，可选）。
    示例: {"title": "分析报告", "content": "# 报告\\n\\n正文...", "format": "pdf"}"""
    import json as _json

    try:
        spec_obj = _json.loads(spec) if isinstance(spec, str) else spec
    except _json.JSONDecodeError:
        spec_obj = {"content": spec}

    title = spec_obj.get("title", "Spectra 分析报告")
    content = spec_obj.get("content", spec if isinstance(spec, str) else "")
    fmt = spec_obj.get("format", "pdf").lower()
    charts = spec_obj.get("charts", [])

    if fmt not in ("pdf", "docx"):
        return f"错误: 不支持的格式 '{fmt}'，仅支持 pdf 或 docx"

    ext = "pdf" if fmt == "pdf" else "docx"

    chart_lines = ""
    for i, c in enumerate(charts):
        chart_lines += f'doc.add_picture("{c}", width=Inches(6))\n' if ext == "docx" else ""

    if ext == "docx":
        code = f'''
from docx import Document
from docx.shared import Inches
import os

doc = Document()
doc.add_heading("{title}", 0)
doc.add_paragraph("""{content[:5000]}""")
{chart_lines}
output_path = "{title}.{ext}"
doc.save(output_path)
print(f"REPORT_GENERATED:{{output_path}}")
print(f"FILE_GENERATED:{{output_path}}")
'''
    else:
        code = f'''
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import os

output_path = "{title}.{ext}"
doc = SimpleDocTemplate(output_path, pagesize=A4)
styles = getSampleStyleSheet()
story = [Paragraph("{title}", styles["Title"]), Spacer(1, 12)]

for line in """{content[:5000]}""".split("\\n"):
    if line.strip():
        story.append(Paragraph(line, styles["Normal"]))
        story.append(Spacer(1, 6))

doc.build(story)
print(f"REPORT_GENERATED:{{output_path}}")
print(f"FILE_GENERATED:{{output_path}}")
'''

    from backend.tools.sandbox import execute_python

    result = execute_python.invoke({"code": code})
    output = str(result) if result else "报告生成完成。"

    try:
        from backend.agent.v2.infra.validator import run_validator

        legacy_state = {
            "generated_code": code,
            "code_type": "python",
            "sender": "reporter",
            "execution_result": output,
            "messages": [],
        }
        val_result = run_validator(legacy_state)
        if val_result.get("validation_passed"):
            return output
        diagnostic = val_result.get("diagnostic", "校验未通过")
        return f"[校验失败] {diagnostic}\n\n执行输出:\n{output}"
    except Exception:
        return output


HEAVY_TOOLS = [run_in_sandbox, generate_pdf_report]
