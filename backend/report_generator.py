"""
报告导出器 —— 把 Markdown 文本 + ECharts 图片（前端 toDataURL 得到的 PNG）
拼成一份可读的 docx / pdf 报告。

设计要点：
- 不再导出对话记录（user/assistant 交替）。直接把 LLM 生成的报告 Markdown
  正文渲染为结构化文档，图表按占位符替换为 PNG。
- 占位符约定：Markdown 中出现 `<agentArtifact type="echarts" title="...">...</agentArtifact>`
  或裸 ECharts JSON 块时，按"出现顺序"绑定到调用方传进来的 charts 列表里
  的 PNG（front-end 直接用 echarts.getDataURL 拿 PNG dataURL）。
- 兼容旧入口：当 content 是 message 数组时，会尝试取最后一条非空 assistant 消息
  作为正文，丢掉对话头尾文字。
"""

from __future__ import annotations

import base64
import io
import os
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable, Optional

from backend.app_paths import ARTIFACTS_DIR, ensure_directories


# ─────────────────────────────────────────────
# 公共入口
# ─────────────────────────────────────────────


def generate_report(
    format_type: str,
    title: str,
    content: Any,
    thread_id: str = "export",
    charts: Optional[list[dict]] = None,
    sources: Optional[list[dict]] = None,
):
    """生成 docx/pdf 报告。

    参数：
        format_type: "pdf" 或 "docx"
        title: 文档主标题（备用，若 markdown 内有 H1 则优先使用 H1）
        content: 报告正文，支持以下三种形式：
            - str: 直接当作 Markdown 正文
            - list[dict]: 旧版 messages 数组，会取最后一条非空 assistant 消息
            - 其他类型: str(...)
        thread_id: 用于子目录名
        charts: 由前端 echarts.getDataURL 收集的图表 PNG，结构：
            [{"chartId": str, "title": str, "dataUrl": "data:image/png;base64,..."}]
        sources: 文末参考资料列表，结构：
            [{"index": int, "title": str, "url": str}]

    返回：
        {"file_path": "/files/<thread>/<filename>", "filename": ..., "absolute_path": ...}
    """
    ensure_directories()

    thread_dir = ARTIFACTS_DIR / thread_id
    thread_dir.mkdir(parents=True, exist_ok=True)

    markdown_text = _coerce_to_markdown(content)
    markdown_text = _clean_markdown_for_export(markdown_text)
    chart_lookup = _normalize_charts(charts)
    sources = sources or []

    # 优先使用 Markdown 第一个 H1 作为标题
    h1_title = _extract_first_h1(markdown_text)
    final_title = h1_title or title or "Spectra 分析报告"

    safe_title = re.sub(r'[\\/*?:"<>|]', '_', final_title)[:80] or "spectra_report"
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    ext = ".pdf" if format_type == "pdf" else ".docx"
    filename = f"{safe_title}_{timestamp}{ext}"
    file_path = thread_dir / filename

    blocks = parse_markdown_blocks(markdown_text, chart_lookup)
    if sources:
        blocks.extend(_build_sources_blocks(sources))

    if format_type == "pdf":
        _write_pdf(blocks, final_title, file_path)
    else:
        _write_docx(blocks, final_title, file_path)

    rel_path = f"/files/{thread_id}/{filename}"
    return {"file_path": rel_path, "filename": filename, "absolute_path": str(file_path)}


# ─────────────────────────────────────────────
# 输入归一化
# ─────────────────────────────────────────────


def _coerce_to_markdown(content: Any) -> str:
    if isinstance(content, str):
        return content
    if isinstance(content, list):
        # 旧入口：messages 数组。取最后一条非空 assistant 消息作为正文。
        for msg in reversed(content):
            if isinstance(msg, dict) and msg.get("role") == "assistant":
                text = msg.get("content") or ""
                if text.strip():
                    return text
        # 全是空消息或无 assistant，退化为拼接所有文字
        return "\n\n".join(
            (m.get("content") or "") for m in content if isinstance(m, dict)
        )
    return str(content)


def _normalize_charts(charts: Optional[list[dict]]) -> dict[str, dict]:
    """把 charts 列表整理成 {chartId: {title, png_bytes}}，并保留顺序索引。"""
    if not charts:
        return {}
    out: dict[str, dict] = {}
    for idx, item in enumerate(charts):
        if not isinstance(item, dict):
            continue
        chart_id = str(item.get("chartId") or item.get("id") or f"_anon_{idx}")
        data_url = item.get("dataUrl") or item.get("data_url") or ""
        png_bytes = _decode_data_url(data_url)
        out[chart_id] = {
            "order": idx,
            "title": item.get("title") or "",
            "png": png_bytes,
        }
    return out


def _decode_data_url(data_url: str) -> Optional[bytes]:
    if not data_url or not isinstance(data_url, str):
        return None
    if "," not in data_url:
        return None
    try:
        b64 = data_url.split(",", 1)[1]
        return base64.b64decode(b64)
    except Exception:
        return None


def _extract_first_h1(text: str) -> str:
    if not text:
        return ""
    for line in text.splitlines():
        m = re.match(r'\s*#\s+(.+?)\s*$', line)
        if m:
            return m.group(1).strip()
    return ""


# ─────────────────────────────────────────────
# 对话残留清洗
# ─────────────────────────────────────────────
#
# Single Agent 给前端的 markdown 经常夹杂两类对话语句：
#   1. 开场白：第一段普通段落里出现"我已为您..."、"下面是..."等聊天式开场
#   2. delivery_suggestion 末尾：SINGLE_AGENT_BASE_PROMPT 要求 LLM 在分析完成
#      时追加一句"是否需要我帮您整理导出为 PDF 或 DOCX 格式的文档？"
# 这些在聊天里有用，但导到正式文档里很违和。导出前剥掉。

# 末尾的"导出引导"句式，常见变体：
#   "是否需要我帮您整理导出为 PDF 或 DOCX 格式的文档？（回复『导出 pdf』...）"
#   "是否需要我将该份分析导出为 PDF 或 DOCX 格式？（直复『导出pdf』或『导出docx』即可）"
#   "要不要我帮您把这份分析导出为 PDF 文档？回复『导出 pdf』即可"
# 启发式：末段同时出现疑问/邀请词 + 导出动词 + 格式名 + 问号，整段剔除
_DELIVERY_SUGGESTION_RE = re.compile(
    r"(?:^|\n)[ \t]*"  # 段落起点
    r"(?=[^\n]*(?:是否|要不要|需不需要|需要|是否需要))"  # 含疑问邀请词
    r"(?=[^\n]*(?:整理|导出|保存|生成|做成|转成|存为|输出))"  # 含导出动词
    r"(?=[^\n]*(?:pdf|docx|word|文档))"  # 含格式名
    r"[^\n]*[?？][^\n]*$",  # 含问号收尾
    re.IGNORECASE,
)

# 开场白：纯粹的聊天过场句子，不包含任何数据/结论关键词。
# 只匹配明确以"我来帮您..."、"下面我给..."等开头的纯服务性语句，且整段不超过80字符。
_CONVERSATIONAL_OPENER_RE = re.compile(
    r"^\s*"
    r"(?:我(?:来帮|已为|会给|将为|这就给|马上给|现在给|先给)(?:您|你))"
    r"[^\n]*$"
)


def _clean_markdown_for_export(text: str) -> str:
    """把 LLM 对话回复转成更适合"正式文档"的 markdown。

    规则（保守，宁少勿多，避免误删正文）：
    1. 删掉末尾的"导出引导"句（delivery_suggestion）
    2. 如果第一段是纯聊天式开场白且后面紧跟 H1/H2 或图表标签，剥掉这一段
    """
    if not text:
        return text

    # 1. 末尾导出引导：从后往前找最后一段，命中模式就删掉这一段
    cleaned = _DELIVERY_SUGGESTION_RE.sub("", text).rstrip()

    # 2. 第一段若是聊天式开场白且紧接结构化内容，删掉
    lines = cleaned.split("\n")
    # 找到第一段非空内容的范围 [first, end)
    first = 0
    while first < len(lines) and not lines[first].strip():
        first += 1
    end = first
    while end < len(lines) and lines[end].strip():
        end += 1
    if first < end:
        first_paragraph = " ".join(l.strip() for l in lines[first:end]).strip()
        # 段落不长 + 命中聊天开场模式 + 段落里没图表标签
        if (
            len(first_paragraph) <= 80
            and _CONVERSATIONAL_OPENER_RE.match(first_paragraph)
            and "<agentArtifact" not in first_paragraph
            and "|" not in first_paragraph  # 不要误伤表格行
        ):
            # 跳过这一段以及它后面的空行
            j = end
            while j < len(lines) and not lines[j].strip():
                j += 1
            cleaned = "\n".join(lines[j:])

    return cleaned.strip()


# ─────────────────────────────────────────────
# Markdown 块解析
# ─────────────────────────────────────────────
#
# 解析输出的 block 形如：
#   {"type": "h1"|"h2"|"h3"|"p"|"ul"|"ol"|"hr"|"image"|"table", ...}
#
# image 块包含 png 字节流（如有）。配对规则：
#   1. 命中 <agentArtifact ...> 标签时，按"出现顺序"配对 charts 列表
#   2. 命中裸 ECharts JSON 块时，同样按顺序消费下一个未用过的 chart
#   3. 找不到对应 PNG 时，渲染成"[图表 N: 标题]"占位段落

_AGENT_ARTIFACT_RE = re.compile(
    r'<agentArtifact\s+type="echarts"\s+title="([^"]*)"\s*>([\s\S]*?)</agentArtifact>',
    re.IGNORECASE,
)


def parse_markdown_blocks(
    text: str,
    chart_lookup: dict[str, dict],
) -> list[dict]:
    """把 markdown 文本切成块序列，把图表占位符替换为 image 块。

    chart_lookup 的图按"出现顺序"消费：第 N 个占位符吃 charts[N]。
    """
    if not text:
        return []

    # 顺序排布的图列表（按 order 排序）
    ordered_charts = sorted(chart_lookup.values(), key=lambda x: x["order"])
    chart_iter = iter(ordered_charts)

    # 替换 <agentArtifact> 与裸 JSON 为唯一的占位符 token，便于后续切分
    PLACEHOLDER = "\u0000__SPECTRA_CHART_{idx}__\u0000"
    placeholders: list[dict] = []  # idx -> {title, png_bytes}

    def _consume_chart(title_hint: str = "") -> tuple[str, Optional[bytes]]:
        try:
            cdata = next(chart_iter)
            return cdata.get("title") or title_hint, cdata.get("png")
        except StopIteration:
            return title_hint, None

    def _replace_artifact(m: re.Match) -> str:
        title = (m.group(1) or "").strip()
        chart_title, png = _consume_chart(title)
        idx = len(placeholders)
        placeholders.append({"title": chart_title or title, "png": png})
        return PLACEHOLDER.format(idx=idx)

    text = _AGENT_ARTIFACT_RE.sub(_replace_artifact, text)

    # 再扫一遍裸 JSON ECharts 块（兜底）
    text = _replace_bare_echarts_blocks(
        text, _consume_chart, placeholders, PLACEHOLDER,
    )

    # 现在按行解析
    blocks: list[dict] = []
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()
        stripped = line.strip()

        # 占位符
        ph_match = re.match(
            r"^\u0000__SPECTRA_CHART_(\d+)__\u0000\s*$", stripped
        )
        if ph_match:
            idx = int(ph_match.group(1))
            if idx < len(placeholders):
                ph = placeholders[idx]
                blocks.append({
                    "type": "image",
                    "title": ph.get("title") or "",
                    "png": ph.get("png"),
                    "fallback_index": idx + 1,
                })
            i += 1
            continue

        # 行内含占位符（混在文字段里），先把占位符所在段拆出来
        if "\u0000__SPECTRA_CHART_" in stripped:
            sub_blocks = _split_inline_placeholder(stripped, placeholders)
            blocks.extend(sub_blocks)
            i += 1
            continue

        # 标题
        h_match = re.match(r"^(#{1,6})\s+(.+?)\s*$", stripped)
        if h_match:
            level = len(h_match.group(1))
            blocks.append({"type": f"h{min(level, 4)}", "text": h_match.group(2)})
            i += 1
            continue

        # 水平分隔
        if re.match(r"^\s*[-*_]{3,}\s*$", line):
            blocks.append({"type": "hr"})
            i += 1
            continue

        # 表格
        if "|" in line and i + 1 < len(lines) and re.match(r"^\s*\|?\s*[-:|\s]+\|?\s*$", lines[i + 1]):
            tbl_lines: list[str] = [line]
            j = i + 1
            while j < len(lines) and "|" in lines[j]:
                tbl_lines.append(lines[j])
                j += 1
            tbl = _parse_table(tbl_lines)
            if tbl:
                blocks.append(tbl)
                i = j
                continue

        # 列表（无序）
        if re.match(r"^\s*[-*+]\s+", line):
            items, consumed = _consume_list(lines, i, ordered=False)
            blocks.append({"type": "ul", "items": items})
            i = consumed
            continue

        # 列表（有序）
        if re.match(r"^\s*\d+[\.\)]\s+", line):
            items, consumed = _consume_list(lines, i, ordered=True)
            blocks.append({"type": "ol", "items": items})
            i = consumed
            continue

        # 空行
        if not stripped:
            i += 1
            continue

        # 普通段落：把后续连续非空、非特殊起始行合并
        para_lines = [stripped]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].rstrip()
            nxt_strip = nxt.strip()
            if not nxt_strip:
                break
            if re.match(r"^(#{1,6})\s+", nxt_strip):
                break
            if re.match(r"^\s*[-*+]\s+", nxt) or re.match(r"^\s*\d+[\.\)]\s+", nxt):
                break
            if "\u0000__SPECTRA_CHART_" in nxt_strip:
                break
            if "|" in nxt and j + 1 < len(lines) and re.match(r"^\s*\|?\s*[-:|\s]+\|?\s*$", lines[j + 1]):
                break
            para_lines.append(nxt_strip)
            j += 1
        blocks.append({"type": "p", "text": " ".join(para_lines)})
        i = j

    return blocks


def _split_inline_placeholder(text: str, placeholders: list[dict]) -> list[dict]:
    """处理同一行里夹杂的占位符。"""
    parts: list[dict] = []
    cursor = 0
    for m in re.finditer(r"\u0000__SPECTRA_CHART_(\d+)__\u0000", text):
        before = text[cursor:m.start()].strip()
        if before:
            parts.append({"type": "p", "text": before})
        idx = int(m.group(1))
        if idx < len(placeholders):
            ph = placeholders[idx]
            parts.append({
                "type": "image",
                "title": ph.get("title") or "",
                "png": ph.get("png"),
                "fallback_index": idx + 1,
            })
        cursor = m.end()
    tail = text[cursor:].strip()
    if tail:
        parts.append({"type": "p", "text": tail})
    return parts


def _consume_list(lines: list[str], start: int, ordered: bool) -> tuple[list[str], int]:
    pat = r"^\s*\d+[\.\)]\s+(.+)$" if ordered else r"^\s*[-*+]\s+(.+)$"
    rx = re.compile(pat)
    items: list[str] = []
    i = start
    while i < len(lines):
        m = rx.match(lines[i])
        if not m:
            break
        items.append(m.group(1).strip())
        i += 1
    return items, i


def _parse_table(tbl_lines: list[str]) -> Optional[dict]:
    if len(tbl_lines) < 2:
        return None

    def split_row(s: str) -> list[str]:
        s = s.strip()
        if s.startswith("|"):
            s = s[1:]
        if s.endswith("|"):
            s = s[:-1]
        return [c.strip() for c in s.split("|")]

    header = split_row(tbl_lines[0])
    rows = [split_row(line) for line in tbl_lines[2:] if line.strip()]
    return {"type": "table", "header": header, "rows": rows}


def _replace_bare_echarts_blocks(
    text: str,
    consume_chart,
    placeholders: list[dict],
    placeholder_template: str,
) -> str:
    """识别裸 ECharts JSON 并替换为占位符（与前端兜底解析逻辑一致）。"""
    out_parts: list[str] = []
    i = 0
    while i < len(text):
        open_idx = text.find("{", i)
        if open_idx == -1:
            out_parts.append(text[i:])
            break
        # 把开括号前的内容原样保留
        out_parts.append(text[i:open_idx])

        end_idx = _match_balanced_json(text, open_idx)
        if end_idx == -1:
            # 未匹配到完整 JSON，原样保留剩余
            out_parts.append(text[open_idx:])
            break

        candidate = text[open_idx:end_idx + 1]
        if _looks_like_echarts(candidate):
            chart_title, png = consume_chart()
            idx = len(placeholders)
            placeholders.append({"title": chart_title, "png": png})
            out_parts.append(placeholder_template.format(idx=idx))
            i = end_idx + 1
            continue

        # 不是 ECharts，原样保留这个 { ... }
        out_parts.append(candidate)
        i = end_idx + 1

    return "".join(out_parts)


def _match_balanced_json(text: str, open_idx: int) -> int:
    depth = 0
    in_str = False
    escape = False
    for k in range(open_idx, len(text)):
        ch = text[k]
        if in_str:
            if escape:
                escape = False
                continue
            if ch == "\\":
                escape = True
                continue
            if ch == '"':
                in_str = False
            continue
        if ch == '"':
            in_str = True
            continue
        if ch == "{":
            depth += 1
        elif ch == "}":
            depth -= 1
            if depth == 0:
                return k
            if depth < 0:
                return -1
    return -1


def _looks_like_echarts(s: str) -> bool:
    if not s or len(s) < 30:
        return False
    if '"series"' not in s:
        return False
    return any(k in s for k in ('"xAxis"', '"yAxis"', '"radar"', '"polar"')) or \
        bool(re.search(r'"type"\s*:\s*"pie"', s))


def _build_sources_blocks(sources: list[dict]) -> list[dict]:
    blocks: list[dict] = [{"type": "h2", "text": "参考资料"}]
    items: list[str] = []
    for src in sources:
        if not isinstance(src, dict):
            continue
        title = (src.get("title") or src.get("url") or "").strip()
        url = (src.get("url") or "").strip()
        if title and url and title != url:
            items.append(f"{title} — {url}")
        elif url:
            items.append(url)
        elif title:
            items.append(title)
    if items:
        blocks.append({"type": "ol", "items": items})
    return blocks


# ─────────────────────────────────────────────
# DOCX 渲染
# ─────────────────────────────────────────────


def _write_docx(blocks: list[dict], title: str, output_path: Path) -> None:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    # 主标题
    head = doc.add_heading(title, level=0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for blk in blocks:
        t = blk.get("type", "p")
        if t == "h1":
            doc.add_heading(blk["text"], level=1)
        elif t == "h2":
            doc.add_heading(blk["text"], level=2)
        elif t == "h3":
            doc.add_heading(blk["text"], level=3)
        elif t == "h4":
            doc.add_heading(blk["text"], level=4)
        elif t == "p":
            p = doc.add_paragraph()
            _docx_render_inline(p, blk["text"])
        elif t == "hr":
            p = doc.add_paragraph("─" * 40)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif t == "ul":
            for it in blk.get("items", []):
                p = doc.add_paragraph(style="List Bullet")
                _docx_render_inline(p, it)
        elif t == "ol":
            for it in blk.get("items", []):
                p = doc.add_paragraph(style="List Number")
                _docx_render_inline(p, it)
        elif t == "table":
            header = blk.get("header") or []
            rows = blk.get("rows") or []
            if header:
                tbl = doc.add_table(rows=1 + len(rows), cols=len(header))
                tbl.style = "Light Grid Accent 1"
                for ci, cell_text in enumerate(header):
                    cell = tbl.rows[0].cells[ci]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)
                    run.bold = True
                for ri, row in enumerate(rows, start=1):
                    for ci, cell_text in enumerate(row[: len(header)]):
                        tbl.rows[ri].cells[ci].text = cell_text
                doc.add_paragraph()  # 表格后空一行
        elif t == "image":
            png = blk.get("png")
            if png:
                try:
                    img_p = doc.add_paragraph()
                    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = img_p.add_run()
                    run.add_picture(io.BytesIO(png), width=Inches(6.0))
                except Exception:
                    doc.add_paragraph(f"[图表 {blk.get('fallback_index', '')}: {blk.get('title', '')}]")
            else:
                doc.add_paragraph(f"[图表 {blk.get('fallback_index', '')}: {blk.get('title', '')}]")
            cap = blk.get("title") or ""
            if cap:
                cap_p = doc.add_paragraph()
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = cap_p.add_run(f"图：{cap}")
                run.italic = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x6B, 0x72, 0x80)

    doc.save(str(output_path))


_INLINE_PATTERN = re.compile(
    r"(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`|\[[^\]]+\]\([^)]+\))"
)


def _docx_render_inline(paragraph, text: str) -> None:
    """简单渲染 markdown 行内格式：**bold**、*italic*、`code`、[label](url)。"""
    if not text:
        return
    cursor = 0
    for m in _INLINE_PATTERN.finditer(text):
        if m.start() > cursor:
            paragraph.add_run(text[cursor:m.start()])
        token = m.group(1)
        if token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("*") and token.endswith("*"):
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
        elif token.startswith("[") and token.endswith(")"):
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mm:
                r = paragraph.add_run(mm.group(1))
                r.font.color.rgb = _docx_blue()
                r.underline = True
        cursor = m.end()
    if cursor < len(text):
        paragraph.add_run(text[cursor:])


def _docx_blue():
    from docx.shared import RGBColor
    return RGBColor(0x25, 0x63, 0xEB)


# ─────────────────────────────────────────────
# PDF 渲染
# ─────────────────────────────────────────────


def _write_pdf(blocks: list[dict], title: str, output_path: Path) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        Image,
        ListFlowable,
        ListItem,
        Paragraph,
        SimpleDocTemplate,
        Spacer,
        Table,
        TableStyle,
    )

    candidates = [
        ("MicrosoftYaHei", r"C:\Windows\Fonts\msyh.ttc"),
        ("MicrosoftYaHeiUI", r"C:\Windows\Fonts\msyh.ttf"),
        ("SimHei", r"C:\Windows\Fonts\simhei.ttf"),
        ("SimSun", r"C:\Windows\Fonts\simsun.ttc"),
    ]
    font_name = "Helvetica"
    for fn, fp in candidates:
        if os.path.exists(fp):
            try:
                pdfmetrics.registerFont(TTFont(fn, fp))
                font_name = fn
                break
            except Exception:
                continue
    if font_name == "Helvetica":
        try:
            pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
            font_name = "STSong-Light"
        except Exception:
            pass

    base = getSampleStyleSheet()
    s_title = ParagraphStyle(
        "ReportTitle", parent=base["Title"], fontName=font_name,
        fontSize=22, leading=28, textColor=colors.HexColor("#1a1a2e"),
        spaceAfter=10,
    )
    s_h1 = ParagraphStyle(
        "ReportH1", parent=base["Heading1"], fontName=font_name,
        fontSize=16, leading=22, textColor=colors.HexColor("#16213e"),
        spaceBefore=12, spaceAfter=6,
    )
    s_h2 = ParagraphStyle(
        "ReportH2", parent=base["Heading2"], fontName=font_name,
        fontSize=13, leading=18, textColor=colors.HexColor("#16213e"),
        spaceBefore=10, spaceAfter=4,
    )
    s_h3 = ParagraphStyle(
        "ReportH3", parent=base["Heading3"], fontName=font_name,
        fontSize=11, leading=15, textColor=colors.HexColor("#374151"),
        spaceBefore=8, spaceAfter=3,
    )
    s_h4 = ParagraphStyle(
        "ReportH4", parent=base["Heading4"], fontName=font_name,
        fontSize=10, leading=14, textColor=colors.HexColor("#4b5563"),
        spaceBefore=6, spaceAfter=2,
    )
    s_body = ParagraphStyle(
        "ReportBody", parent=base["BodyText"], fontName=font_name,
        fontSize=10, leading=16, textColor=colors.HexColor("#1f2937"),
    )
    s_caption = ParagraphStyle(
        "ReportCaption", parent=base["BodyText"], fontName=font_name,
        fontSize=8.5, leading=12, textColor=colors.HexColor("#6b7280"),
        alignment=1,  # CENTER
    )

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=A4,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
    )
    story: list = [Paragraph(title, s_title), Spacer(1, 6 * mm)]

    page_width = A4[0] - 40 * mm

    for blk in blocks:
        t = blk.get("type", "p")
        if t == "h1":
            story.append(Paragraph(_pdf_inline(blk["text"]), s_h1))
        elif t == "h2":
            story.append(Paragraph(_pdf_inline(blk["text"]), s_h2))
        elif t == "h3":
            story.append(Paragraph(_pdf_inline(blk["text"]), s_h3))
        elif t == "h4":
            story.append(Paragraph(_pdf_inline(blk["text"]), s_h4))
        elif t == "p":
            story.append(Paragraph(_pdf_inline(blk["text"]), s_body))
            story.append(Spacer(1, 4))
        elif t == "hr":
            story.append(Spacer(1, 6))
        elif t == "ul":
            items = [
                ListItem(Paragraph(_pdf_inline(it), s_body), leftIndent=14)
                for it in blk.get("items", [])
            ]
            story.append(ListFlowable(items, bulletType="bullet", start="•", leftIndent=14))
            story.append(Spacer(1, 4))
        elif t == "ol":
            items = [
                ListItem(Paragraph(_pdf_inline(it), s_body), leftIndent=14)
                for it in blk.get("items", [])
            ]
            story.append(ListFlowable(items, bulletType="1", leftIndent=14))
            story.append(Spacer(1, 4))
        elif t == "table":
            header = blk.get("header") or []
            rows = blk.get("rows") or []
            if header:
                data = [[Paragraph(_pdf_inline(c), s_body) for c in header]]
                for r in rows:
                    data.append([Paragraph(_pdf_inline(c), s_body) for c in r[: len(header)]])
                tbl = Table(data, hAlign="LEFT", repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f3f4f6")),
                    ("FONTNAME", (0, 0), (-1, -1), font_name),
                    ("FONTSIZE", (0, 0), (-1, -1), 9),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d1d5db")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(tbl)
                story.append(Spacer(1, 6))
        elif t == "image":
            png = blk.get("png")
            if png:
                try:
                    img = Image(io.BytesIO(png))
                    iw, ih = img.imageWidth, img.imageHeight
                    if iw and ih:
                        scale = min(page_width / iw, 1.0)
                        img.drawWidth = iw * scale
                        img.drawHeight = ih * scale
                    story.append(img)
                except Exception:
                    story.append(Paragraph(
                        f"[图表 {blk.get('fallback_index', '')}: {blk.get('title', '')}]",
                        s_body,
                    ))
            else:
                story.append(Paragraph(
                    f"[图表 {blk.get('fallback_index', '')}: {blk.get('title', '')}]",
                    s_body,
                ))
            cap = blk.get("title") or ""
            if cap:
                story.append(Paragraph(f"图：{cap}", s_caption))
            story.append(Spacer(1, 6))

    doc.build(story)


_PDF_INLINE_PATTERN = re.compile(
    r"(\*\*[^*\n]+?\*\*|\*[^*\n]+?\*|`[^`\n]+?`|\[[^\]]+\]\([^)]+\))"
)


def _pdf_inline(text: str) -> str:
    """把简单 markdown 行内格式转为 reportlab 支持的 mini HTML。"""
    if not text:
        return ""
    # reportlab 只识别有限的实体；先转义关键字符
    safe = (
        text.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
    )

    def _repl(m: re.Match) -> str:
        token = m.group(1)
        # 注意：因为我们已经做过 HTML escape，原始 markdown token 仍能匹配
        if token.startswith("**") and token.endswith("**"):
            return f"<b>{token[2:-2]}</b>"
        if token.startswith("*") and token.endswith("*"):
            return f"<i>{token[1:-1]}</i>"
        if token.startswith("`") and token.endswith("`"):
            return f'<font face="Courier">{token[1:-1]}</font>'
        if token.startswith("[") and token.endswith(")"):
            mm = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if mm:
                return f'<link href="{mm.group(2)}" color="#2563eb">{mm.group(1)}</link>'
        return token

    return _PDF_INLINE_PATTERN.sub(_repl, safe)
