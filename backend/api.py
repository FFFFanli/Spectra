import os
import json
import re
import uuid
import base64
from datetime import datetime
from pathlib import Path
from fastapi import FastAPI, UploadFile, File, Request, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import HTMLResponse, FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from starlette.middleware.base import BaseHTTPMiddleware
from sse_starlette.sse import EventSourceResponse
import uvicorn
from apscheduler.schedulers.asyncio import AsyncIOScheduler
from apscheduler.triggers.cron import CronTrigger

# 启动时优先加载项目根目录 .env，确保 E2B/模型等配置可直接生效
try:
    from dotenv import load_dotenv

    env_path = Path(__file__).resolve().parent.parent / ".env"
    if env_path.exists():
        load_dotenv(env_path)
except Exception:
    pass

from backend.app_paths import ARTIFACTS_DIR, DIST_DIR, FRONTEND_DIR, STATIC_DIR, ensure_directories
from backend.db_utils import save_file_to_duckdb

from backend.agent.single_agent import build_single_agent_graph
from backend.agent.v2.runtime import TeamOrchestrationRuntime
from backend.tools import ALL_TOOLS as AGENT_TOOLS
from backend.tools.sandbox import SandboxSession
from backend.tools.task_manager import set_thread_id as set_gtd_thread_id
from backend.request_context import begin_request, get_usage_summary, get_request_model, set_attached_charts, set_export_content, set_last_assistant_reply
from backend.skill_loader import find_skill_for_tool
from langchain_core.messages import HumanMessage, AIMessage
from backend.state_store import (
    add_alert,
    init_state_store,
    list_alerts,
)
from backend.conversation_store import (
    init_conversation_store,
    list_conversations,
    get_conversation,
    upsert_conversation,
    delete_conversation,
    clear_conversations,
    DEFAULT_USER_ID,
)
from backend.checkpoint_store import init_checkpoint_store, close_checkpoint_store
from backend.memory import retrieve_memory_context, save_structured_memory
from backend.tools.user_memory import set_memory_user_id
from backend.tools.cron_manager import set_scheduler, _run_cron_task
from backend.tools.user_interaction import get_pending_question, clear_pending_question

# ───────────── Access Code 鉴权 ─────────────
_ACCESS_CODE = os.environ.get("SPECTRA_ACCESS_CODE", "").strip()

class AccessCodeMiddleware(BaseHTTPMiddleware):
    async def dispatch(self, request: Request, call_next):
        if _ACCESS_CODE and request.url.path.startswith("/api"):
            auth = request.headers.get("Authorization", "")
            if not auth:
                return JSONResponse(
                    status_code=401,
                    content={"error": "未提供 Access Code，请在 Authorization header 中传入 Bearer <code>"},
                )
            token = auth.removeprefix("Bearer ").strip()
            if token != _ACCESS_CODE:
                return JSONResponse(
                    status_code=401,
                    content={"error": "Access Code 错误"},
                )
        return await call_next(request)

app = FastAPI()

scheduler = AsyncIOScheduler()

@app.on_event("startup")
async def start_scheduler():
    ensure_directories()
    init_state_store()
    init_conversation_store()
    await init_checkpoint_store()

    # 将 scheduler 注入 cron 工具模块
    set_scheduler(scheduler)

    # 从数据库恢复所有 active 的 cron 任务
    from backend.state_store import get_active_cron_jobs
    active_jobs = get_active_cron_jobs()
    restored = 0
    for job in active_jobs:
        try:
            trigger = CronTrigger.from_crontab(job["cron_expr"])
            scheduler.add_job(
                _run_cron_task,
                trigger,
                args=[job["prompt"]],
                id=job["job_id"],
            )
            restored += 1
        except Exception as e:
            print(f"[Cron] 恢复任务失败 {job['job_id']}: {e}")
    if restored > 0:
        print(f"[Cron] 已从数据库恢复 {restored} 个定时任务")

    scheduler.start()


@app.on_event("shutdown")
async def stop_scheduler():
    scheduler.shutdown()
    await close_checkpoint_store()

# 允许跨域
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)

# Access Code 鉴权
app.add_middleware(AccessCodeMiddleware)

# 挂载产物目录以便下载文件和预览图表
app.mount("/files", StaticFiles(directory=str(ARTIFACTS_DIR)), name="files")
# 挂载本地的 CDN 静态资源目录
app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")
# 挂载 Vite 构建产物 (生产模式)
dist_assets = DIST_DIR / "assets"
if dist_assets.exists():
    app.mount("/assets", StaticFiles(directory=str(dist_assets)), name="assets")

@app.post("/api/settings")
async def update_settings(request: Request):
    """更新用户的 API Keys 配置"""
    data = await request.json()
    dashscope_key = data.get("dashscope_key", "").strip()
    openai_key = data.get("openai_key", "").strip()
    deepseek_key = data.get("deepseek_key", "").strip()
    selected_model = data.get("selected_model", "").strip()
    
    if dashscope_key:
        os.environ["DASHSCOPE_API_KEY"] = dashscope_key
    if openai_key:
        os.environ["OPENAI_API_KEY"] = openai_key
    if deepseek_key:
        os.environ["DEEPSEEK_API_KEY"] = deepseek_key
    if selected_model:
        os.environ["SPECTRA_SELECTED_MODEL"] = selected_model
        
    return {"status": "ok"}

# ───────────── 模型列表 ─────────────
_MODEL_REGISTRY = {
    "qwen3.5-plus": {"name": "通义千问 Qwen3.5-Plus", "provider": "dashscope"},
    "qwen3.6-plus": {"name": "通义千问 Qwen3.6-Plus", "provider": "dashscope"},
    "qwen-max": {"name": "通义千问 Qwen-Max", "provider": "dashscope"},
    "gpt-4o": {"name": "OpenAI GPT-4o", "provider": "openai"},
    "deepseek-v4-flash": {"name": "DeepSeek V4 Flash", "provider": "deepseek"},
}

def _get_available_models() -> list[dict]:
    """根据已配置的 API Key 动态返回可用模型列表"""
    has_dashscope = bool(os.environ.get("DASHSCOPE_API_KEY"))
    has_openai = bool(os.environ.get("OPENAI_API_KEY"))
    has_deepseek = bool(os.environ.get("DEEPSEEK_API_KEY"))
    models = []
    for model_id, info in _MODEL_REGISTRY.items():
        provider = info["provider"]
        if provider == "dashscope" and has_dashscope:
            models.append({"id": model_id, "name": info["name"]})
        elif provider == "openai" and has_openai:
            models.append({"id": model_id, "name": info["name"]})
        elif provider == "deepseek" and has_deepseek:
            models.append({"id": model_id, "name": info["name"]})
    # 如果一个 key 都没有，返回全部（允许用户在设置页填入 key 后刷新）
    if not models:
        models = [{"id": mid, "name": info["name"]} for mid, info in _MODEL_REGISTRY.items()]
    return models

@app.get("/api/models")
async def list_models():
    """返回当前可用的模型列表"""
    return {"models": _get_available_models()}

@app.post("/api/upload")
async def upload_file(file: UploadFile = File(...)):
    """处理文件上传，存入 DuckDB"""
    if not file.filename:
        raise HTTPException(status_code=400, detail="缺少文件名")

    lower_name = file.filename.lower()
    if lower_name.endswith((".csv", ".xlsx", ".xls")):
        import tempfile
        temp_path = os.path.join(tempfile.gettempdir(), f"upload_{file.filename}")
        with open(temp_path, "wb") as f:
            f.write(await file.read())

        with open(temp_path, "rb") as f:
            df, table_name = save_file_to_duckdb(f, file.filename)
        return {"table_name": table_name, "rows": len(df), "file_type": "table"}

    if lower_name.endswith((".pdf", ".json")):
        ensure_directories()
        suffix = Path(file.filename).suffix
        target = ARTIFACTS_DIR / f"upload_{uuid.uuid4().hex[:8]}{suffix}"
        with open(target, "wb") as f:
            f.write(await file.read())
        return {
            "file_type": "pdf_template" if lower_name.endswith(".pdf") else "json_context",
            "path": target.name,
            "filename": file.filename,
        }

    raise HTTPException(status_code=400, detail="仅支持 .csv、.xlsx、.xls、.pdf 和 .json 文件")

@app.post("/api/connect_db")
async def connect_db(request: Request):
    """处理外部数据库直连请求"""
    data = await request.json()
    db_type = data.get("db_type")
    connection_string = data.get("connection_string")
    alias = data.get("alias", "ext_db")

    if not db_type or not connection_string:
        raise HTTPException(status_code=400, detail="缺少必填参数 db_type 或 connection_string")

    try:
        from backend.db_utils import attach_external_database
        await asyncio.to_thread(attach_external_database, db_type, connection_string, alias)
        return {"status": "success", "message": f"成功连接外部数据库 {alias}"}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/profile")
async def profile(request: Request):
    """自动生成数据洞察体检报告"""
    data = await request.json()
    table_name = data.get("table_name")

    if not table_name:
        raise HTTPException(status_code=400, detail="缺少表名参数")

    try:
        from backend.db_utils import generate_data_profile
        result = await asyncio.to_thread(generate_data_profile, table_name)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/tables")
async def list_tables():
    """获取 DuckDB 中所有用户表名"""
    try:
        from backend.db_utils import get_all_tables
        tables = await asyncio.to_thread(get_all_tables)
        return {"tables": tables}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/table_data/{table_name}")
async def table_data(table_name: str, limit: int = 100):
    """获取指定表的前 N 行数据"""
    try:
        from backend.db_utils import get_table_preview
        result = await asyncio.to_thread(get_table_preview, table_name, limit)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

import asyncio


@app.post("/api/chat")
async def chat(request: Request):
    """统一 Agent SSE 端点 —— Solo Agent + 完整工具面板。"""
    data = await request.json()
    user_message = data.get("message", "") or ""
    original_query = user_message
    persona_system_prompt = data.get("persona_system_prompt", "").strip()

    # 前端传来的 system_prompt 作为用户自定义追加（不再包含 base prompt）
    user_extra_prompt = data.get("system_prompt", "").strip()
    if user_extra_prompt == "你是一个智能助手，可以使用联网搜索和网页爬取工具来获取最新信息，帮助用户解决问题。":
        user_extra_prompt = ""  # 前端默认值，等同未传

    # 检索历史相关记忆（结构化：preference / fact / experience / context）
    user_id = _resolve_user_id(request)
    set_memory_user_id(user_id)
    memory_context = await asyncio.to_thread(retrieve_memory_context, user_id, user_message)
    if memory_context:
        user_extra_prompt = f"{memory_context}\n\n{user_extra_prompt}".strip()

    # 附带图表处理
    attached_charts_input = data.get("attached_charts") or []
    decoded_charts: list[dict] = []
    for idx, item in enumerate(attached_charts_input):
        if not isinstance(item, dict):
            continue
        data_url = item.get("dataUrl") or item.get("data_url") or ""
        if not isinstance(data_url, str) or "," not in data_url:
            continue
        try:
            png_bytes = base64.b64decode(data_url.split(",", 1)[1])
        except Exception:
            continue
        if not png_bytes:
            continue
        suffix = ".png"
        raw_name = (item.get("name") or "").strip()
        if raw_name and raw_name.lower().endswith((".png", ".jpg", ".jpeg")):
            file_name = raw_name
        else:
            file_name = f"chart_{idx + 1}{suffix}"
        decoded_charts.append({
            "name": file_name,
            "title": item.get("title") or "",
            "png_bytes": png_bytes,
        })

    attached_charts_brief = ""
    if decoded_charts:
        lines = ["", "<attached_charts>", "**重要：用户消息已附带前端已渲染好的图表 PNG，沙盒里这些文件已就绪：**"]
        for c in decoded_charts:
            title = c["title"] or "（无标题）"
            lines.append(f"- `{c['name']}`：{title}")
        lines.extend([
            "",
            "当用户要求导出 docx/pdf 时，**直接在 python-docx / reportlab 中读取这些 PNG 嵌入文档**，",
            "不要用 matplotlib / plotly 重画。示例：",
            "```python",
            "from docx import Document",
            "from docx.shared import Inches",
            "doc = Document()",
            "doc.add_heading('分析报告', 0)",
            "doc.add_paragraph('正文...')",
            "doc.add_picture('chart_1.png', width=Inches(6))",
            "doc.save('report.docx')",
            "print('FILE_GENERATED:report.docx')",
            "```",
            "</attached_charts>",
        ])
        attached_charts_brief = "\n".join(lines)

    # 附件透传：用户上传文件对应的 DuckDB 表 + 外部数据库别名
    attached_files = data.get("attached_files") or []
    db_alias = data.get("db_alias", "").strip()

    available_data_brief = ""
    if attached_files or db_alias:
        lines = ["", "<available_data>", "**以下数据表可供直接查询（已入库 DuckDB）：**"]
        for f in attached_files:
            table = f.get("table_name", "")
            name = f.get("name", "")
            ftype = f.get("type", "csv")
            if table:
                lines.append(f"- `{table}`：{name}（{ftype}）")
        if db_alias:
            lines.append(f"- 外部数据库别名：`{db_alias}`（通过 DuckDB ATTACH 连接，用 `query_duckdb(\"SELECT * FROM {db_alias}.表名\")` 查询）")
        lines.append("</available_data>")
        available_data_brief = "\n".join(lines)

    export_hint = ""
    export_content = data.get("export_content") or {}
    is_export_request = _looks_like_export(user_message)
    if is_export_request:
        fmt = "pdf" if "pdf" in user_message.lower() else "docx"
        if export_content and (export_content.get("content") or "").strip():
            # 前端已注入图表 PNG + markdown → generate_docx 直接从 ContextVar 取
            export_hint = (
                "\n\n<export_request>\n"
                f"**这是一次导出请求。请直接调用 generate_docx(format=\"{fmt}\")，不要传 content 参数。**\n"
                "前端已注入完整 markdown 和图表 PNG，工具会自动取用。\n"
                "**禁止**调 execute_python / run_in_sandbox 自己写代码生成文档，那条路径慢且丢图。\n"
                "</export_request>"
            )
        else:
            # 前端没预注入内容 → LLM 需要自己把最近回复的 markdown 传给 generate_docx
            export_hint = (
                "\n\n<export_request>\n"
                f"**这是一次导出请求。请调用 generate_docx(format=\"{fmt}\")，"
                "并把你在上一条回复中生成的完整 markdown 正文（保留所有 <agentArtifact> 标签）"
                "作为 content 参数传入。**\n"
                "**禁止**调 execute_python / run_in_sandbox 自己写代码生成文档，那条路径慢且丢图。\n"
                "</export_request>"
            )

    from backend.prompt_loader import build_system_prompt as assemble_system_prompt

    system_prompt = assemble_system_prompt(
        persona_prompt=persona_system_prompt,
        user_prompt=user_extra_prompt,
        export_hint=export_hint,
        attached_charts_brief=attached_charts_brief,
        available_data_brief=available_data_brief,
    )
    thread_id = data.get("thread_id", str(uuid.uuid4()))
    model_name = data.get("model", "qwen3.6-plus")
    begin_request(model_name)

    if decoded_charts:
        set_attached_charts(decoded_charts)
    if export_content:
        set_export_content(export_content)

    single_graph = build_single_agent_graph(
        tools=AGENT_TOOLS,
        system_prompt=system_prompt,
        max_steps=40,
    )

    # 将 thread_id 注入 GTD 任务工具，确保任务按对话线程隔离
    set_gtd_thread_id(thread_id)

    async def agent_event_generator():
        final_reply = ""
        emitted_text = False
        try:
            with SandboxSession():
                async for event in single_graph.astream_events(
                    {"messages": [HumanMessage(content=user_message)], "system_prompt": system_prompt},
                    config={"configurable": {"thread_id": thread_id}, "recursion_limit": 100},
                    version="v2",
                ):
                    kind = event["event"]

                    if kind == "on_chat_model_stream":
                        chunk = event["data"]["chunk"]
                        rc = (getattr(chunk, "additional_kwargs", {}) or {}).get("reasoning_content")
                        if rc:
                            yield {
                                "event": "reasoning_stream",
                                "data": json.dumps({"content": rc}, ensure_ascii=False),
                            }
                        if chunk.content:
                            final_reply += str(chunk.content)
                            emitted_text = True
                            yield {
                                "event": "llm_stream",
                                "data": json.dumps({"content": chunk.content}, ensure_ascii=False),
                            }

                    elif kind == "on_chat_model_end":
                        # 兜底：如果流式过程 final_reply 为空（chunk.content 全空），
                        # 但最终消息有 content，补发 llm_stream 事件
                        output = event["data"].get("output")
                        if (
                            output is not None
                            and hasattr(output, "content")
                            and output.content
                            and not final_reply
                        ):
                            text = str(output.content)
                            final_reply = text
                            emitted_text = True
                            yield {
                                "event": "llm_stream",
                                "data": json.dumps({"content": text}, ensure_ascii=False),
                            }

                    elif kind == "on_tool_start":
                        tool_name = event["name"]
                        skill = find_skill_for_tool(tool_name)
                        payload = {
                            "tool": tool_name,
                            "input": event["data"].get("input", {}),
                        }
                        if skill:
                            payload["skill"] = {
                                "name": skill.name,
                                "capability": skill.capability,
                            }
                        yield {
                            "event": "tool_start",
                            "data": json.dumps(payload, ensure_ascii=False),
                        }

                    elif kind == "on_tool_end":
                        raw_output = event["data"].get("output", "")
                        if hasattr(raw_output, "content"):
                            tool_output = str(raw_output.content or "")
                        else:
                            tool_output = str(raw_output)
                        yield {
                            "event": "tool_result",
                            "data": json.dumps({
                                "tool": event["name"],
                                "output": tool_output[:500],
                            }, ensure_ascii=False),
                        }
                        if event["name"] == "export_conversation" and "EXPORT_FILE:" in tool_output:
                            match = re.search(r"EXPORT_FILE:([^|]+)\|([^|]+)\|(\w+)", tool_output)
                            if match:
                                yield {
                                    "event": "file",
                                    "data": json.dumps({
                                        "url": match.group(1),
                                        "name": match.group(2),
                                        "format": match.group(3),
                                    }, ensure_ascii=False),
                                }
                        if event["name"] in ("execute_python", "generate_docx", "run_in_sandbox", "generate_pdf_report"):
                            for evt in _parse_sandbox_artifact_markers(tool_output):
                                yield evt

                        # 检测用户交互工具调用 → 发射 user_question 事件
                        if event["name"] in ("ask_user", "request_confirmation"):
                            question = get_pending_question()
                            if question:
                                yield {
                                    "event": "user_question",
                                    "data": json.dumps(question, ensure_ascii=False),
                                }
                                clear_pending_question()

            usage_summary = get_usage_summary()
            yield {"event": "usage", "data": json.dumps(usage_summary, ensure_ascii=False)}

            if final_reply.strip():
                asyncio.create_task(asyncio.to_thread(
                    save_structured_memory,
                    user_id,
                    f"用户需求: {original_query[:300]}\n回复摘要: {final_reply[:300]}",
                    "context",
                    thread_id,
                    "chat",
                ))
                set_last_assistant_reply(final_reply)

            # [Spectra debug] 诊断日志
            print(f"[Spectra SSE] done — final_reply len={len(final_reply)}, thread={thread_id[:8]}", flush=True)

            # 兜底：某些模型/SDK 可能在内部拿到了最终文本，但前端没有成功收到正文片段。
            # done 前补发一条完整 reply，避免界面出现"后端有答案但前端空白"。
            if final_reply.strip() and not emitted_text:
                yield {
                    "event": "reply",
                    "data": json.dumps(final_reply, ensure_ascii=False),
                }

            yield {"event": "done", "data": json.dumps({"thread_id": thread_id})}

        except Exception as e:
            import traceback
            traceback.print_exc()
            yield {"event": "error", "data": str(e)}

    headers = {
        "X-Accel-Buffering": "no",
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
    }

    return EventSourceResponse(agent_event_generator(), headers=headers)


@app.post("/api/v2/chat")
async def chat_v2(request: Request):
    """Team Supervisor v2 SSE 端点 —— 多 Agent 编排模式。"""
    data = await request.json()
    user_message = data.get("message", "") or ""
    original_query = user_message
    persona_system_prompt = data.get("persona_system_prompt", "").strip()
    user_extra_prompt = data.get("system_prompt", "").strip()
    if user_extra_prompt == "你是一个智能助手，可以使用联网搜索和网页爬取工具来获取最新信息，帮助用户解决问题。":
        user_extra_prompt = ""

    # 检索记忆
    user_id = _resolve_user_id(request)
    set_memory_user_id(user_id)
    memory_context = await asyncio.to_thread(retrieve_memory_context, user_id, user_message)
    if memory_context:
        user_extra_prompt = f"{memory_context}\n\n{user_extra_prompt}".strip()

    # 附带图表处理
    attached_charts_input = data.get("attached_charts") or []
    decoded_charts: list[dict] = []
    for idx, item in enumerate(attached_charts_input):
        if not isinstance(item, dict):
            continue
        data_url = item.get("dataUrl") or item.get("data_url") or ""
        if not isinstance(data_url, str) or "," not in data_url:
            continue
        try:
            png_bytes = base64.b64decode(data_url.split(",", 1)[1])
        except Exception:
            continue
        if not png_bytes:
            continue
        suffix = ".png"
        raw_name = (item.get("name") or "").strip()
        if raw_name and raw_name.lower().endswith((".png", ".jpg", ".jpeg")):
            file_name = raw_name
        else:
            file_name = f"chart_{idx + 1}{suffix}"
        decoded_charts.append({
            "name": file_name,
            "title": item.get("title") or "",
            "png_bytes": png_bytes,
        })

    # 附件透传：数据表信息注入 schema
    attached_files = data.get("attached_files") or []
    db_alias = data.get("db_alias", "").strip()
    schema_lines = []
    if attached_files:
        for f in attached_files:
            table = f.get("table_name", "")
            name = f.get("name", "")
            if table:
                schema_lines.append(f"表 `{table}`：{name}")
    if db_alias:
        schema_lines.append(f"外部数据库：`{db_alias}`")
    schema = "\n".join(schema_lines) if schema_lines else ""

    # 导出请求处理
    export_content = data.get("export_content") or {}
    is_export_request = _looks_like_export(user_message)
    if is_export_request and export_content:
        set_export_content(export_content)

    thread_id = data.get("thread_id", str(uuid.uuid4()))
    model_name = data.get("model", "qwen3.6-plus")
    begin_request(model_name)

    if decoded_charts:
        set_attached_charts(decoded_charts)

    runtime = TeamOrchestrationRuntime()

    async def v2_event_generator():
        final_reply = ""
        try:
            with SandboxSession():
                async for event in runtime.run(
                    user_message=user_message,
                    thread_id=thread_id,
                    schema=schema,
                    conversation_history=None,
                ):
                    event_type = event.get("event", "")
                    event_data = event.get("data", {})

                    if event_type == "supervisor_decision":
                        yield {
                            "event": "supervisor_decision",
                            "data": json.dumps(event_data, ensure_ascii=False),
                        }

                    elif event_type == "agent_message":
                        agent_id = event_data.get("agent_id", "")
                        reply = event_data.get("reply", "")
                        code = event_data.get("code", "")
                        yield {
                            "event": "agent_message",
                            "data": json.dumps({
                                "agent_id": agent_id,
                                "reply": reply[:500],
                                "has_code": bool(code),
                            }, ensure_ascii=False),
                        }

                    elif event_type == "reply":
                        final_reply = event_data.get("text", "")
                        yield {
                            "event": "reply",
                            "data": json.dumps({"text": final_reply}, ensure_ascii=False),
                        }

                    elif event_type == "done":
                        yield {
                            "event": "done",
                            "data": json.dumps({**event_data, "thread_id": thread_id}, ensure_ascii=False),
                        }

                    elif event_type == "error":
                        yield {
                            "event": "error",
                            "data": json.dumps({"message": event_data.get("message", str(event_data))}, ensure_ascii=False),
                        }

            usage_summary = get_usage_summary()
            yield {"event": "usage", "data": json.dumps(usage_summary, ensure_ascii=False)}

            if final_reply.strip():
                asyncio.create_task(asyncio.to_thread(
                    save_structured_memory,
                    user_id,
                    f"用户需求: {original_query[:300]}\n回复摘要: {final_reply[:300]}",
                    "context",
                    thread_id,
                    "v2_chat",
                ))
                set_last_assistant_reply(final_reply)

            print(f"[Spectra SSE v2] done — final_reply len={len(final_reply)}, thread={thread_id[:8]}", flush=True)

        except Exception as e:
            import traceback
            traceback.print_exc()
            yield {"event": "error", "data": str(e)}

    headers = {
        "X-Accel-Buffering": "no",
        "Cache-Control": "no-cache",
        "Connection": "keep-alive",
    }

    return EventSourceResponse(v2_event_generator(), headers=headers)


def _looks_like_export(msg: str) -> bool:
    """检测用户消息是否包含导出文档意图。"""
    if not msg or not isinstance(msg, str):
        return False
    verbs = "导出|保存|存为|存成|另存|转成|转为|整理成|打成|生成|下载成|做成"
    formats = "pdf|docx|word|文档"
    import re as _re
    return bool(
        _re.search(verbs, msg) and _re.search(formats, msg, _re.IGNORECASE)
    )


def _parse_sandbox_artifact_markers(tool_output: str) -> list[dict]:
    """从 execute_python 工具的 stdout 中解析沙盒产物 marker，转为 SSE 事件。

    沙盒已把文件搬到 ARTIFACTS_DIR，stdout 里的路径已被改写为本地相对路径。
    这里把 marker 转成前端能消费的 file / artifacts 事件。

    Marker 约定（与 backend/tools/sandbox.py 的 _harvest_e2b_artifacts 对齐）：
      CHART_GENERATED:<rel>       → 图表 HTML
      CHART_PNG_GENERATED:<rel>   → 图表 PNG
      CLEANED_DATA_GENERATED:<rel>→ 清洗后的数据文件
      REPORT_GENERATED:<rel>      → PDF/DOCX 报告
      FILE_GENERATED:<rel>        → 通用产物（docx / xlsx / 等）
    """
    if not tool_output:
        return []

    events: list[dict] = []
    artifacts_payload: list[dict] = []

    rules: list[tuple[str, str, str]] = [
        # (marker, default_format, artifact_type)
        ("CHART_GENERATED:", "HTML", "chart_html"),
        ("CHART_PNG_GENERATED:", "PNG", "chart_png"),
        ("CLEANED_DATA_GENERATED:", "XLSX", "cleaned_data"),
        ("REPORT_GENERATED:", "REPORT", "report"),
        ("FILE_GENERATED:", "FILE", "file"),
    ]

    for marker, default_format, item_type in rules:
        for raw_path in re.findall(rf"{re.escape(marker)}([^\r\n]+)", tool_output):
            rel = raw_path.strip()
            if not rel:
                continue
            url = f"/files/{rel.lstrip('/')}"
            name = rel.rsplit("/", 1)[-1]
            suffix = name.rsplit(".", 1)[-1].upper() if "." in name else default_format
            actual_type = item_type
            if marker == "REPORT_GENERATED:":
                lower = name.lower()
                if lower.endswith(".pdf"):
                    actual_type = "report_pdf"
                elif lower.endswith(".docx"):
                    actual_type = "report_docx"
            artifacts_payload.append({
                "type": actual_type,
                "name": name,
                "path": url,
            })
            # 对可下载文档（docx/pdf/xlsx 等）下发 file 事件，前端会渲染下载卡片
            if actual_type in {"report_pdf", "report_docx", "file", "cleaned_data"} or marker == "FILE_GENERATED:":
                events.append({
                    "event": "file",
                    "data": json.dumps({
                        "url": url,
                        "name": name,
                        "format": suffix,
                    }, ensure_ascii=False),
                })

    if artifacts_payload:
        # 复用现有 artifacts 通道，前端会自动加进 taskArtifacts / charts 列表
        events.insert(0, {
            "event": "artifacts",
            "data": json.dumps(artifacts_payload, ensure_ascii=False),
        })

    return events

@app.post("/api/export_conversation")
async def export_conversation(request: Request):
    data = await request.json()
    thread_id = data.get("thread_id", "unknown")
    format_type = data.get("format", "pdf").lower()
    title = data.get("title", "Spectra 分析报告")

    # 新版：前端会传 report_markdown + charts；旧版兼容 messages
    report_markdown = data.get("report_markdown")
    charts = data.get("charts") or []
    sources = data.get("sources") or []
    messages = data.get("messages", [])

    if format_type not in ("pdf", "docx"):
        raise HTTPException(status_code=400, detail="不支持的导出格式，仅支持 pdf 或 docx")

    if not report_markdown and not messages:
        raise HTTPException(status_code=400, detail="没有可导出的内容")

    content = report_markdown if report_markdown else messages

    try:
        result = await asyncio.to_thread(
            generate_report,
            format_type,
            title,
            content,
            thread_id,
            charts,
            sources,
        )
        return result
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

from backend.report_generator import generate_report


@app.get("/api/workflows")
async def list_workflows():
    """列出所有预定义工作流模板（Phase 1 后已精简为统一 Agent 模式）。"""
    return {"workflows": []}


def _run_automated_solo_task(prompt: str):
    """定时巡检后台任务：用 Solo Agent + 统一工具面板跑全流程，结果写入 alerts。"""
    from backend.agent.prompts import CHART_PROMPT, SANDBOX_SYSTEM_PROMPT

    begin_request("")
    thread_id = str(uuid.uuid4())

    system_prompt = (
        SANDBOX_SYSTEM_PROMPT
        + "\n\n你是一个智能助手，可以使用联网搜索和网页爬取工具来获取最新信息。"
        + CHART_PROMPT
    )

    graph = build_single_agent_graph(
        tools=AGENT_TOOLS,
        system_prompt=system_prompt,
        max_steps=40,
    )

    final_reply = ""
    chart_paths: list[str] = []

    async def run_and_collect() -> tuple[str, list[str]]:
        nonlocal final_reply, chart_paths
        with SandboxSession():
            async for event in graph.astream_events(
                {"messages": [HumanMessage(content=prompt)], "system_prompt": system_prompt},
                config={"configurable": {"thread_id": thread_id}, "recursion_limit": 100},
                version="v2",
            ):
                kind = event["event"]
                if kind == "on_chat_model_stream":
                    chunk = event["data"]["chunk"]
                    if chunk.content:
                        final_reply += str(chunk.content)
                elif kind == "on_tool_end":
                    raw_output = event["data"].get("output", "")
                    if hasattr(raw_output, "content"):
                        tool_output = str(raw_output.content or "")
                    else:
                        tool_output = str(raw_output)
                    for marker in ("CHART_GENERATED:", "CHART_PNG_GENERATED:"):
                        for raw_path in re.findall(rf"{re.escape(marker)}([^\r\n]+)", tool_output):
                            chart_paths.append(f"/files/{raw_path.strip().lstrip('/')}")
        return final_reply or "巡检任务已完成，未生成具体报告文本。", chart_paths

    try:
        report_text, charts = asyncio.run(run_and_collect())
        add_alert(
            alert_id=str(uuid.uuid4()),
            prompt=prompt,
            report=report_text,
            charts=charts,
        )
        print(f"[Alert] 巡检任务执行成功: {prompt}")
    except Exception as e:
        import traceback
        traceback.print_exc()
        add_alert(
            alert_id=str(uuid.uuid4()),
            prompt=prompt,
            report=f"⚠️ 巡检任务执行失败: {str(e)}",
            charts=[],
        )

@app.post("/api/schedule")
async def create_schedule(request: Request):
    """创建定时巡检任务（前端直接调用）"""
    data = await request.json()
    prompt = data.get("prompt")
    cron_expr = data.get("cron")

    if not prompt or not cron_expr:
        raise HTTPException(status_code=400, detail="缺少 prompt 或 cron")

    try:
        trigger = CronTrigger.from_crontab(cron_expr)
        job_id = f"cron_{uuid.uuid4().hex[:8]}"
        scheduler.add_job(_run_cron_task, trigger, args=[prompt], id=job_id)
        # 同时持久化到数据库
        from backend.state_store import save_cron_job
        save_cron_job(job_id, cron_expr, prompt)
        return {"status": "success", "message": "定时任务已创建", "job_id": job_id}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"创建定时任务失败: {str(e)}")

@app.get("/api/alerts")
async def get_alerts():
    """获取所有巡检预警记录"""
    return {"alerts": list_alerts()}


# ---------------- 对话历史持久化 ----------------
# 跨浏览器/换机同步对话记录。当前默认所有用户共享 DEFAULT_USER_ID，
# 接入登录后把 _resolve_user_id 替换为从请求里取真实身份即可。

def _resolve_user_id(request: Request) -> str:
    # 预留：未来从 cookie / header 中读取真实用户身份
    return request.headers.get("x-user-id") or DEFAULT_USER_ID


@app.get("/api/conversations")
async def api_list_conversations(request: Request):
    """获取当前用户的对话历史列表（仅元数据）"""
    user_id = _resolve_user_id(request)
    items = await asyncio.to_thread(list_conversations, user_id)
    return {"items": items}


@app.get("/api/conversations/{conv_id}")
async def api_get_conversation(conv_id: str, request: Request):
    """获取单个对话的完整内容"""
    user_id = _resolve_user_id(request)
    item = await asyncio.to_thread(get_conversation, conv_id, user_id)
    if not item:
        raise HTTPException(status_code=404, detail="对话不存在")
    return item


@app.post("/api/conversations/{conv_id}")
async def api_save_conversation(conv_id: str, request: Request):
    """新增或覆盖一条对话"""
    user_id = _resolve_user_id(request)
    data = await request.json()
    thread_id = (data.get("threadId") or data.get("thread_id") or "").strip()
    title = (data.get("title") or "空对话").strip()
    messages = data.get("messages") or []
    if not isinstance(messages, list):
        raise HTTPException(status_code=400, detail="messages 必须是数组")
    meta = await asyncio.to_thread(
        upsert_conversation,
        conv_id=conv_id,
        thread_id=thread_id,
        title=title,
        messages=messages,
        user_id=user_id,
    )
    return meta


@app.delete("/api/conversations/{conv_id}")
async def api_delete_conversation(conv_id: str, request: Request):
    """删除单条对话"""
    user_id = _resolve_user_id(request)
    ok = await asyncio.to_thread(delete_conversation, conv_id, user_id)
    return {"ok": ok}


@app.delete("/api/conversations")
async def api_clear_conversations(request: Request):
    """清空当前用户全部对话"""
    user_id = _resolve_user_id(request)
    deleted = await asyncio.to_thread(clear_conversations, user_id)
    return {"deleted": deleted}


@app.get("/")
async def get_index():
    """返回前端页面 —— 优先使用 Vite 构建产物，否则回退到源码 index.html (dev 模式)"""
    dist_html = DIST_DIR / "index.html"
    if dist_html.exists():
        return FileResponse(dist_html)
    return FileResponse(FRONTEND_DIR / "index.html")

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
